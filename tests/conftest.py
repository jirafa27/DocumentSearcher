import asyncio
import io
import shutil
import tempfile
from pathlib import Path

import pytest
import pytest_asyncio
from docx import Document
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker

from app.core.config import settings
from app.core.database import Base, get_async_session
from app.main import app

# Переключаемся на auto режим для pytest-asyncio
pytest_plugins = ("pytest_asyncio",)


@pytest.fixture(scope="session")
def event_loop():
    """Создает event loop для async тестов"""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest_asyncio.fixture
async def test_db_session():
    """Простая фикстура для тестовой БД - создаем/удаляем схему для каждого теста"""
    engine = create_async_engine(
        "postgresql+asyncpg://postgres:postgres@localhost:5433/test_documentsearcher",
        echo=False,
    )

    # Создаем схему для этого теста
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)

    # Создаем сессию
    async_session = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

    async with async_session() as session:
        yield session

    # Удаляем схему после теста
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)

    await engine.dispose()


@pytest.fixture
def test_upload_dir():
    """Создает временную папку для тестовых загрузок"""
    temp_dir = tempfile.mkdtemp(prefix="test_uploads_")
    original_upload_dir = settings.UPLOAD_DIR

    # Переопределяем папку загрузок для тестов
    settings.UPLOAD_DIR = temp_dir

    yield temp_dir

    # Восстанавливаем оригинальную настройку и удаляем временную папку
    settings.UPLOAD_DIR = original_upload_dir
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest_asyncio.fixture
async def test_client(test_db_session, test_upload_dir):
    """Тестовый HTTP клиент с переопределенными зависимостями"""

    # Переопределяем зависимость БД для тестов
    def override_get_async_session():
        yield test_db_session

    app.dependency_overrides[get_async_session] = override_get_async_session

    # Новый синтаксис для httpx AsyncClient
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        yield client

    app.dependency_overrides.clear()


@pytest.fixture
def db_checker(test_db_session):
    """Утилита для проверки записей в БД"""
    from sqlalchemy import select

    from app.models.document import Document as SQLDocument
    from app.models.document import DocumentContent as SQLDocumentContent

    async def get_document_by_id(document_id):
        """Получить документ из БД по ID"""
        query = select(SQLDocument).where(SQLDocument.id == document_id)
        result = await test_db_session.execute(query)
        return result.scalar_one_or_none()

    async def get_document_content_by_document_id(document_id):
        """Получить содержимое документа из БД по ID документа"""
        query = select(SQLDocumentContent).where(
            SQLDocumentContent.document_id == document_id
        )
        result = await test_db_session.execute(query)
        return result.scalar_one_or_none()

    async def count_documents():
        """Подсчитать количество документов в БД"""
        query = select(SQLDocument)
        result = await test_db_session.execute(query)
        return len(result.scalars().all())

    async def count_document_contents():
        """Подсчитать количество записей содержимого в БД"""
        query = select(SQLDocumentContent)
        result = await test_db_session.execute(query)
        return len(result.scalars().all())

    return {
        "get_document": get_document_by_id,
        "get_content": get_document_content_by_document_id,
        "count_documents": count_documents,
        "count_contents": count_document_contents,
        "session": test_db_session,
    }


@pytest.fixture
def sample_file():
    """Базовая фикстура с тестовым текстовым файлом"""
    return {
        "filename": "test.txt",
        "content": b"Test document content for integration testing",
        "content_type": "text/plain",
    }


@pytest.fixture
def sample_file_docx():
    """Простой DOCX файл для тестирования"""

    doc = Document()
    doc.add_paragraph("Test document content for integration testing")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return {
        "filename": "test.docx",
        "content": buffer.getvalue(),
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


@pytest.fixture
def sample_file_pdf():
    """Фикстура с настоящим PDF документом"""
    import io

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()

    # Создаем PDF документ
    pdf = canvas.Canvas(buffer, pagesize=letter)

    # Добавляем текст (используем стандартный шрифт для простоты)
    pdf.drawString(100, 750, "Test document content for integration testing")
    pdf.drawString(100, 730, "Отдел продаж компании показал отличные результаты")

    # Завершаем создание PDF
    pdf.save()

    # Получаем содержимое
    buffer.seek(0)

    return {
        "filename": "test.pdf",
        "content": buffer.getvalue(),
        "content_type": "application/pdf",
    }


@pytest.fixture
def sample_file_large():
    """Фикстура с большим документом для тестирования производительности"""
    large_content = b"Large document content. " * 1000  # ~24KB
    return {
        "filename": "large_test.txt",
        "content": large_content,
        "content_type": "text/plain",
    }


@pytest.fixture
def invalid_file_format():
    """Фикстура с файлом неподдерживаемого формата"""
    return {
        "filename": "test.xyz",
        "content": b"Invalid file content with unsupported format",
        "content_type": "application/unknown",
    }


@pytest.fixture
def invalid_file_too_large():
    """Фикстура с файлом, превышающим максимальный размер"""
    # Создаем файл больше 20MB (лимит в настройках)
    large_content = b"X" * (21 * 1024 * 1024)  # 21MB
    return {
        "filename": "huge_file.pdf",
        "content": large_content,
        "content_type": "application/pdf",
    }


@pytest.fixture
def search_test_document():
    """
    Специальный документ для тестирования поиска
    Содержит разные словоформы, точные фразы и контекст
    """
    content = """
    Отчет о продажах компании за 2024 год

    Отдел продаж показал отличные результаты в этом году.
    Продажи выросли на 25% по сравнению с прошлым годом.
    Продавцы активно работали с клиентами.

    Анализ продаж:
    - Продажа товаров категории A увеличилась
    - Продавать стало проще благодаря новой системе
    - Продавец месяца - Иван Иванов

    Точная фраза для тестирования: "специальный поисковый запрос"

    Контекст для проверки выделения текста и границ фрагментов.
    Этот текст должен помочь протестировать размеры контекста до и после найденных совпадений.

    Дополнительные ключевые слова: документ, тестирование, поиск, система, анализ.
    """

    # Создаем DOCX документ
    doc = Document()
    doc.add_paragraph(content.strip())

    # Сохраняем в байты
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return {
        "filename": "search_test.docx",
        "content": buffer.getvalue(),
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


@pytest.fixture
def uploaded_files_checker(test_upload_dir):
    """Утилита для проверки загруженных файлов в тестах"""

    def check_files_exist(*filenames):
        """Проверяет, что файлы существуют в тестовой папке"""
        upload_path = Path(test_upload_dir)
        existing_files = list(upload_path.glob("*"))
        return len(existing_files) >= len(filenames)

    def get_uploaded_files():
        """Возвращает список загруженных файлов"""
        upload_path = Path(test_upload_dir)
        return list(upload_path.glob("*"))

    def cleanup_files():
        """Очищает все файлы из тестовой папки"""
        upload_path = Path(test_upload_dir)
        for file_path in upload_path.glob("*"):
            file_path.unlink()

    return {
        "check_exist": check_files_exist,
        "get_files": get_uploaded_files,
        "cleanup": cleanup_files,
    }
