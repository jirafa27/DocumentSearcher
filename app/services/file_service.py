import asyncio
import hashlib
import os
from pathlib import Path

import aiofiles
import pdfplumber
from docx import Document as DocxDocument

from app.core.config import settings
from app.core.exceptions.file import (
    FileError,
    FileTooLargeError,
    TextExtractionError,
    UnsupportedFileTypeError,
)
from app.core.interfaces.file_service import IFileService
from app.core.logger import logger


class FileService(IFileService):
    """Сервис для работы с файлами"""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        self.max_file_size = settings.MAX_FILE_SIZE
        self.allowed_types = settings.ALLOWED_FILE_TYPES

    async def validate_file(self, filename: str, content: bytes) -> None:
        """
        Валидация файла. Проверяется тип файла и размер файла.

        Args:
            filename: str - имя файла
            content: bytes - содержимое файла

        Raises:
            UnsupportedFileTypeError: Если тип файла не поддерживается
            FileTooLargeError: Если файл слишком большой
        """
        # Проверка типа файла
        file_extension = filename.split(".")[-1].lower() if filename else ""
        if file_extension not in self.allowed_types:
            logger.error(f"Неподдерживаемый тип файла: {file_extension}")
            raise UnsupportedFileTypeError(self.allowed_types)

        # Проверка размера файла
        file_size = len(content)
        if file_size > self.max_file_size:
            logger.error(
                f"Файл слишком большой. Максимальный размер: {self.max_file_size / 1024 / 1024:.1f}MB"
            )
            raise FileTooLargeError(self.max_file_size)

    async def save_file(self, filename: str, content: bytes, file_hash: str) -> str:
        """
        Сохранение файла в файловую систему

        Args:
            filename: str - оригинальное имя файла
            content: bytes - содержимое файла
            file_hash: str - SHA-256 хеш файла

        Returns:
            str: Путь к сохраненному файлу

        Raises:
            FileError: Если не удалось сохранить файл на диск
        """
        file_path = f"{self.upload_dir}/{file_hash}_{filename}"

        try:
            async with aiofiles.open(file_path, "wb") as f:
                await f.write(content)

            return file_path
        except Exception as e:
            logger.error(f"Ошибка при сохранении файла {filename}: {e}")
            raise FileError(f"Не удалось сохранить файл: {str(e)}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Извлечение текста из PDF файла

        Args:
            file_path: str - путь к файлу

        Returns:
            str: Текст из файла

        Raises:
            TextExtractionError: Если не удалось извлечь текст из PDF
        """
        text = ""

        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise TextExtractionError(file_path, str(e))

        return text.strip()

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Извлечение текста из DOCX файла

        Args:
            file_path: str - путь к файлу

        Returns:
            str: Текст из файла

        Raises:
            TextExtractionError: Если не удалось извлечь текст из DOCX
        """
        try:
            doc = DocxDocument(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"Не удалось извлечь текст из DOCX: {str(e)}")
            raise TextExtractionError(file_path, str(e))

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Извлечение текста в зависимости от типа файла

        Args:
            file_path: str - путь к файлу
            file_type: str - тип файла

        Returns:
            str: Текст из файла

        Raises:
            UnsupportedFileTypeError: Если тип файла не поддерживается
            TextExtractionError: Если не удалось извлечь текст из файла
        """

        await self.validate_file(file_path, file_type)

        if file_type == "pdf":
            # Выполняем в отдельном потоке чтобы не блокировать event loop
            return await asyncio.get_event_loop().run_in_executor(
                None, self._extract_text_from_pdf, file_path
            )
        elif file_type == "docx":
            # Выполняем в отдельном потоке чтобы не блокировать event loop
            return await asyncio.get_event_loop().run_in_executor(
                None, self._extract_text_from_docx, file_path
            )


    async def delete_file(self, file_path: str) -> None:
        """
        Удаление файла с диска

        Args:
            file_path: str - путь к файлу

        Returns:
            None

        Raises:
            FileError: Если не удалось удалить файл
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            else:
                logger.error(f"Файл не найден: {file_path}")
                raise FileNotFoundError(file_path)
        except Exception as e:
            logger.error(f"Не удалось удалить файл: {file_path}")
            raise FileError(f"Не удалось удалить файл: {str(e)}")

    async def calculate_hash(self, content: bytes) -> str:
        """
        Вычисление SHA-256 хеша файла

        Args:
            content: bytes - содержимое файла для вычисления хеша

        Returns:
            str: SHA-256 хеш файла в hex формате
        """
        sha256_hash = hashlib.sha256()
        sha256_hash.update(content)
        hash_result = sha256_hash.hexdigest()
        return hash_result
